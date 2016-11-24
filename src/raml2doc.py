#############################
#
#    copyright 2016 Open Interconnect Consortium, Inc. All rights reserved.
#    Redistribution and use in source and binary forms, with or without modification, are permitted provided that the following conditions are met:
#    1.  Redistributions of source code must retain the above copyright notice, this list of conditions and the following disclaimer.
#    2.  Redistributions in binary form must reproduce the above copyright notice, this list of conditions and the following disclaimer in the documentation and/or other materials provided with the distribution.
#         
#    THIS SOFTWARE IS PROVIDED BY THE OPEN INTERCONNECT CONSORTIUM, INC. "AS IS" AND ANY EXPRESS OR IMPLIED WARRANTIES,
#    INCLUDING, BUT NOT LIMITED TO, THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE OR WARRANTIES OF NON-INFRINGEMENT,
#    ARE DISCLAIMED. IN NO EVENT SHALL THE OPEN INTERCONNECT CONSORTIUM, INC. OR CONTRIBUTORS BE LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY,
#    OR CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS;
#    OR BUSINESS INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
#    OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#
#############################

try:
    import pyramloic.parser as ramlparser
except ImportError:
    import pyraml.parser as ramlparser

    pass
#
# generic imports
#
import os
import sys
import traceback
import argparse
#
# docx imports
#
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# fix for py2exe
from jsonschema import _utils
import json
from yaml import load, dump

try:
    from yaml import CLoader as Loader, CDumper as Dumper
except ImportError:
    from yaml import Loader, Dumper, SafeDumper

draft3schemafile = """{
    "$schema": "http://json-schema.org/draft-03/schema#",
    "dependencies": {
        "exclusiveMaximum": "maximum",
        "exclusiveMinimum": "minimum"
    },
    "id": "http://json-schema.org/draft-03/schema#",
    "properties": {
        "$ref": {
            "format": "uri",
            "type": "string"
        },
        "$schema": {
            "format": "uri",
            "type": "string"
        },
        "additionalItems": {
            "default": {},
            "type": [
                {
                    "$ref": "#"
                },
                "boolean"
            ]
        },
        "additionalProperties": {
            "default": {},
            "type": [
                {
                    "$ref": "#"
                },
                "boolean"
            ]
        },
        "default": {
            "type": "any"
        },
        "dependencies": {
            "additionalProperties": {
                "items": {
                    "type": "string"
                },
                "type": [
                    "string",
                    "array",
                    {
                        "$ref": "#"
                    }
                ]
            },
            "default": {},
            "type": [
                "string",
                "array",
                "object"
            ]
        },
        "description": {
            "type": "string"
        },
        "disallow": {
            "items": {
                "type": [
                    "string",
                    {
                        "$ref": "#"
                    }
                ]
            },
            "type": [
                "string",
                "array"
            ],
            "uniqueItems": true
        },
        "divisibleBy": {
            "default": 1,
            "exclusiveMinimum": true,
            "minimum": 0,
            "type": "number"
        },
        "enum": {
            "minItems": 1,
            "type": "array",
            "uniqueItems": true
        },
        "exclusiveMaximum": {
            "default": false,
            "type": "boolean"
        },
        "exclusiveMinimum": {
            "default": false,
            "type": "boolean"
        },
        "extends": {
            "default": {},
            "items": {
                "$ref": "#"
            },
            "type": [
                {
                    "$ref": "#"
                },
                "array"
            ]
        },
        "format": {
            "type": "string"
        },
        "id": {
            "format": "uri",
            "type": "string"
        },
        "items": {
            "default": {},
            "items": {
                "$ref": "#"
            },
            "type": [
                {
                    "$ref": "#"
                },
                "array"
            ]
        },
        "maxDecimal": {
            "minimum": 0,
            "type": "number"
        },
        "maxItems": {
            "minimum": 0,
            "type": "integer"
        },
        "maxLength": {
            "type": "integer"
        },
        "maximum": {
            "type": "number"
        },
        "minItems": {
            "default": 0,
            "minimum": 0,
            "type": "integer"
        },
        "minLength": {
            "default": 0,
            "minimum": 0,
            "type": "integer"
        },
        "minimum": {
            "type": "number"
        },
        "pattern": {
            "format": "regex",
            "type": "string"
        },
        "patternProperties": {
            "additionalProperties": {
                "$ref": "#"
            },
            "default": {},
            "type": "object"
        },
        "properties": {
            "additionalProperties": {
                "$ref": "#",
                "type": "object"
            },
            "default": {},
            "type": "object"
        },
        "required": {
            "default": false,
            "type": "boolean"
        },
        "title": {
            "type": "string"
        },
        "type": {
            "default": "any",
            "items": {
                "type": [
                    "string",
                    {
                        "$ref": "#"
                    }
                ]
            },
            "type": [
                "string",
                "array"
            ],
            "uniqueItems": true
        },
        "uniqueItems": {
            "default": false,
            "type": "boolean"
        }
    },
    "type": "object"
}
"""

draft4schemafile = """{
    "$schema": "http://json-schema.org/draft-04/schema#",
    "default": {},
    "definitions": {
        "positiveInteger": {
            "minimum": 0,
            "type": "integer"
        },
        "positiveIntegerDefault0": {
            "allOf": [
                {
                    "$ref": "#/definitions/positiveInteger"
                },
                {
                    "default": 0
                }
            ]
        },
        "schemaArray": {
            "items": {
                "$ref": "#"
            },
            "minItems": 1,
            "type": "array"
        },
        "simpleTypes": {
            "enum": [
                "array",
                "boolean",
                "integer",
                "null",
                "number",
                "object",
                "string"
            ]
        },
        "stringArray": {
            "items": {
                "type": "string"
            },
            "minItems": 1,
            "type": "array",
            "uniqueItems": true
        }
    },
    "dependencies": {
        "exclusiveMaximum": [
            "maximum"
        ],
        "exclusiveMinimum": [
            "minimum"
        ]
    },
    "description": "Core schema meta-schema",
    "id": "http://json-schema.org/draft-04/schema#",
    "properties": {
        "$schema": {
            "format": "uri",
            "type": "string"
        },
        "additionalItems": {
            "anyOf": [
                {
                    "type": "boolean"
                },
                {
                    "$ref": "#"
                }
            ],
            "default": {}
        },
        "additionalProperties": {
            "anyOf": [
                {
                    "type": "boolean"
                },
                {
                    "$ref": "#"
                }
            ],
            "default": {}
        },
        "allOf": {
            "$ref": "#/definitions/schemaArray"
        },
        "anyOf": {
            "$ref": "#/definitions/schemaArray"
        },
        "default": {},
        "definitions": {
            "additionalProperties": {
                "$ref": "#"
            },
            "default": {},
            "type": "object"
        },
        "dependencies": {
            "additionalProperties": {
                "anyOf": [
                    {
                        "$ref": "#"
                    },
                    {
                        "$ref": "#/definitions/stringArray"
                    }
                ]
            },
            "type": "object"
        },
        "description": {
            "type": "string"
        },
        "enum": {
            "minItems": 1,
            "type": "array",
            "uniqueItems": true
        },
        "exclusiveMaximum": {
            "default": false,
            "type": "boolean"
        },
        "exclusiveMinimum": {
            "default": false,
            "type": "boolean"
        },
        "format": {
            "type": "string"
        },
        "id": {
            "format": "uri",
            "type": "string"
        },
        "items": {
            "anyOf": [
                {
                    "$ref": "#"
                },
                {
                    "$ref": "#/definitions/schemaArray"
                }
            ],
            "default": {}
        },
        "maxItems": {
            "$ref": "#/definitions/positiveInteger"
        },
        "maxLength": {
            "$ref": "#/definitions/positiveInteger"
        },
        "maxProperties": {
            "$ref": "#/definitions/positiveInteger"
        },
        "maximum": {
            "type": "number"
        },
        "minItems": {
            "$ref": "#/definitions/positiveIntegerDefault0"
        },
        "minLength": {
            "$ref": "#/definitions/positiveIntegerDefault0"
        },
        "minProperties": {
            "$ref": "#/definitions/positiveIntegerDefault0"
        },
        "minimum": {
            "type": "number"
        },
        "multipleOf": {
            "exclusiveMinimum": true,
            "minimum": 0,
            "type": "number"
        },
        "not": {
            "$ref": "#"
        },
        "oneOf": {
            "$ref": "#/definitions/schemaArray"
        },
        "pattern": {
            "format": "regex",
            "type": "string"
        },
        "patternProperties": {
            "additionalProperties": {
                "$ref": "#"
            },
            "default": {},
            "type": "object"
        },
        "properties": {
            "additionalProperties": {
                "$ref": "#"
            },
            "default": {},
            "type": "object"
        },
        "required": {
            "$ref": "#/definitions/stringArray"
        },
        "title": {
            "type": "string"
        },
        "type": {
            "anyOf": [
                {
                    "$ref": "#/definitions/simpleTypes"
                },
                {
                    "items": {
                        "$ref": "#/definitions/simpleTypes"
                    },
                    "minItems": 1,
                    "type": "array",
                    "uniqueItems": true
                }
            ]
        },
        "uniqueItems": {
            "default": false,
            "type": "boolean"
        }
    },
    "type": "object"
}
"""

try:
    from jsonschema import Draft4Validator
    from jsonschema import ValidationError
except:
    # os.mkdir('./jsonschema/schema/')
    f = open("draft3.json", "w")
    f.write(draft3schemafile)
    f.close()
    f = open("draft4.json", "w")
    f.write(draft4schemafile)
    f.close()
    from jsonschema import Draft4Validator


def find_key(rec_dict, target, depth=0):
    try:
        #print (depth,target, rec_dict)
        if isinstance(rec_dict, dict):
            for key,value in rec_dict.items():
                if key == target:
                    return rec_dict[key]
            for key,value in rec_dict.items():
                r = find_key(value, target, depth+1)
                if r is not None:
                        return r
        #else:
        #    print ("no dict:", rec_dict)
    except:
        traceback.print_exc()

# find an key recursively
# also traverse lists (arrays, oneOf,..) but only returns the first occurance
#
#
def find_key_link(rec_dict, target, depth=0):
    if isinstance(rec_dict, dict):
        # direct key
        for key,value in rec_dict.items():
            if key == target:
                return rec_dict[key]
        # key is in array
        rvalues = []
        found = False
        for key,value in rec_dict.items():
            if key in ["oneOf", "allOf", "anyOf"]:
                for val in value:
                    #print ("xxx", depth, key, val)
                    if val == target:
                        return val
                    if isinstance(val, dict):
                        r = find_key_link(val, target, depth+1)
                        if r is not None:
                            found = True
                            #print ("found:", r)
                            # TODO: this should return an array, now it only returns the last found item
                            rvalues = r
                            #rvalues.extend(list(r.items()))
        if found:
            #print (rvalues)
            return rvalues

        # key is an dict
        for key,value in rec_dict.items():
            r = find_key_link(value, target, depth+1)
            if r is not None:
                return r #[list(r.items())]


class CreateDoc(object):
    def __init__(self, name, docx_name=None, resource_name=None):
        """
        initialize the class


        """
        self.annex_switch = False
        self.composite_switch = False
        self.sensor_switch = False
        self.table_method = False
        self.sensor_switch = False
        self.schema_switch = False
        self.schemaWT_switch = False
        self.schemaWT_files = None
        self.schema_files = None
        self.table = None
        self.title = None

        self.inputname = name
        self.resourcedoc = "ResourceTemplate.docx"
        if docx_name is not None:
            if os.path.isfile(docx_name):
                self.resourcedoc = docx_name
            else:
                print "WARNING: could not find file:", docx_name
                print "using:", self.resourcedoc
        self.resource_out = name + ".docx"
        self.tab = "  "
        self.resource_name = resource_name
        self.schema_ignorelist = ['required', '$schema', 'type', 'definitions', 'description',
                                  'properties', ":", ":{", "minItems", "attribute", "format", "allOf", "$ref", "enum",
                                  "title", "oneOf", "anyOf", "additionalProperties", "items", "default", "minitems",
                                  "maxitems",
                                  "minimum", "maximum", "pattern", "readOnly", "minProperties", "additionalItems"]
        self.schema_types = ['boolean', 'array', 'object', 'enum', 'number', 'string']

    def list_resource(self, level, lt_resource, lt_obj):
        """
        function to list the CRUDN behavior per resource
        e.g. it adds an entry to the CRUDN table
        :param level:
        :param resource:
        :param obj:
        :return:
        """
        if lt_obj is None:
            return

        row_cells = self.table.add_row().cells
        # row_cells[0].text = resource
        row_cells[0].text = lt_resource
        if lt_obj.methods is not None:
            for method, mobj in lt_obj.methods.items():
                # print "Method:",method
                # PUT == Create
                if method == "put":
                    row_cells[1].text = method
                # GET = Read
                if method == "get":
                    row_cells[2].text = method
                # POST - update  (agreed on 05/02/2015)
                if method == "post":
                    row_cells[3].text = method
                # DELETE = Delete
                if method == "delete":
                    row_cells[4].text = method
                # NOTIFY = NOTIFY (does not exist)
                if method == "notify":
                    row_cells[5].text = method
        for res_name, res_obj in lt_obj.resources.items():
            self.list_resource(level + 1, res_name, res_obj)

    def list_resources_crudn(self, parse_tree, select_resource=None):
        # function to create the CRUDN table
        """

        :param parse_tree:
        :param select_resource:
        """
        level = 0
        # create the table
        self.table = self.document.add_table(rows=1, cols=6, style='TABLE-A')
        hdr_cells = self.table.rows[0].cells
        hdr_cells[0].text = 'Resource'
        hdr_cells[1].text = 'Create'
        hdr_cells[2].text = 'Read'
        hdr_cells[3].text = 'Update'
        hdr_cells[4].text = 'Delete'
        hdr_cells[5].text = 'Notify'

        if select_resource is None:
            # all resources
            for my_resource, my_obj in parse_tree.resources.items():
                self.list_resource(level, my_resource, my_obj)
        else:
            for my_resource, my_obj in parse_tree.resources.items():
                # only the one of the command line
                if my_resource[1:] == select_resource:
                    self.list_resource(level, my_resource, my_obj)

    def list_description(self, level, resource, obj, select_resource=None):
        """

        :param level:
        :param resource:
        :param obj:
        :param select_resource:
        :return:
        """
        if obj is None:
            return

        if obj is not None:
            if obj.description is not None:
                intro_text = self.remove_eof_smart(obj.description)
                self.document.add_paragraph(intro_text)
        try:
            for nResName, nObj in obj.resources.items():
                if select_resource is None:
                    self.list_description(level + 1, nResName, nObj, select_resource)
                else:
                    if nResName[1:] == select_resource:
                        self.list_description(level + 1, nResName, nObj, select_resource)
        except:
            pass

    def list_descriptions(self, parse_tree, select_resource=None):
        """

        :param parse_tree:
        :param select_resource:
        """
        level = 0

        if select_resource is None:
            for resource, obj in parse_tree.resources.items():
                self.list_description(level, resource, obj)
        else:
            for resource, obj in parse_tree.resources.items():
                if select_resource == resource[1:]:
                    self.list_description(level, resource, obj)

    def list_uri(self, level, resource, obj):

        """

        :param level:
        :param resource:
        :param obj:
        """
        if resource is not None:
            self.document.add_paragraph(resource)

        try:
            for nResName, nObj in obj.resources.items():
                self.list_uri(level + 1, nResName, nObj)
        except:
            pass

    def list_URIs(self, parse_tree, select_resource=None):
        """

        :param parse_tree:
        :param select_resource:
        """
        level = 0

        if select_resource is None:
            for resource, obj in parse_tree.resources.items():
                self.list_uri(level, resource, obj)
        else:
            for resource, obj in parse_tree.resources.items():
                if select_resource == resource[1:]:
                    self.list_uri(level, resource, obj)

    def list_x_resource(self, level, resource, obj, select_resource=None):
        """
        list the resource
        :param level:
        :param resource:
        :param obj:
        :param select_resource:
        :return:
        """
        if obj is None:
            return

        try:
            for nResName, nObj in obj.resources.items():
                self.list_x_resource(level + 1, nResName, nObj, select_resource)
        except:
            pass

    def list_x_resources(self, parse_tree):
        """
        list all resources (loop over all them)
        :param parse_tree:
        """
        for resource, obj in parse_tree.resources.items():
            self.list_x_resource(0, resource, obj)

    def get_display_name_resources(self, parse_tree, resource_name):
        """
        retrieve the display name of the resource
        :param parsetree:
        :param resourceName:
        :return:
        """
        for resource, obj in parse_tree.resources.items():
            # ignore the start slash
            if resource[1:] == resource_name:
                return obj.displayName

    def get_resource_type_line(self, input_lines):
        """
        retrieve rt from the json example
        :param input_lines:
        :return:
        """
        my_input_lines = input_lines.replace(" ", "")
        lines = my_input_lines.splitlines()
        for line in lines:
            tokens = line.split('"')
            if len(tokens) >= 4:
                if tokens[1] == "rt":
                    return tokens[3]
        return None

    def get_resource_type_by_resources(self, parse_tree, resource_name):
        """
        find an example in any body.
        :param parse_tree:
        :param resource_name:
        :return: resource type of the resource name
        """
        for resource, obj in parse_tree.resources.items():
            if resource[1:] == resource_name:
                for method, mobj in obj.methods.items():
                    if mobj.responses is not None:
                        for resName, res in mobj.responses.items():
                            for sName, _body in res.body.items():
                                if sName == "application/json":
                                    value = self.get_resource_type_line(_body.example)
                                    if value is not None:
                                        return value
                                    else:
                                        print "get_resource_type_by_resources ERROR: no RT found in:", _body.example
        return None

    def parse_schema_requires(self, input_string_schema):
        """
        find the required property list

        :param input_string_schema: json schema as string
        :return:
        """
        ignore_list = ['required', '[', ']', ',', ': [']
        lines_schema = input_string_schema.splitlines()
        length = len(lines_schema)
        required_properties = []
        for x in range(0, length - 1):
            # parse a line in a schema
            tokens = lines_schema[x].split('"')
            if len(tokens) > 1:
                if tokens[1] == 'required':
                    for token in tokens:
                        if token == "]":
                            print "correct end of required detected"
                        if token not in ignore_list:
                            if " " not in token:
                                required_properties.append(token)
        return required_properties

    def parse_schema(self, input_string_schema):
        """

        :param input_string_schema:
        """
        required_props = self.parse_schema_requires(input_string_schema)
        print "parse_schema: required properties found:", required_props
        json_dict =json.loads(input_string_schema)

        properties = find_key_link(json_dict, 'properties')

        for prop in properties:
            # fill the table
            try:
                if isinstance(properties, dict):
                    print "parse_schema: property:", prop
                    description_text = properties[prop].get('description', "")
                    read_only = properties[prop].get('readOnly', False)
                    type = properties[prop].get('type')
                    if type is None:
                        type = "multiple types: see schema"
                    if type == "array":
                        type += ": see schema"
                    if type == "object":
                        type += ": see schema"
                    row_cells = self.tableAttribute.add_row().cells
                    row_cells[0].text = str(prop)
                    row_cells[1].text = str(type)
                    if str(prop) in required_props:
                        row_cells[2].text = "yes"
                    if read_only is True:
                        row_cells[3].text = "Read Only"
                    else:
                        row_cells[3].text = "Read Write"
                    row_cells[4].text = description_text

            except:
                traceback.print_exc()
                pass
                
    def list_attribute(self, level, resource, obj):
        """
        list all attributes of an indicated resource
        e.g. put it in the table
        :param level:
        :param resource:
        :param obj:
        :return:
        """
        if obj is None:
            print "EMPTY EMPTY"
            return

        if level != 0:
            if obj.methods is not None:
                for method, mobj in obj.methods.items():
                    # default method get...
                    if method == self.table_method:
                        for res_name, res in mobj.responses.items():
                            if res_name == 200:
                                # we only want the succesfull case
                                for response_type, body in res.body.items():
                                    if response_type == "application/json":
                                        text = self.get_schema_string_from_body(body)
                                        if text is not None:
                                            self.parse_schema(text)

        for res_name, res_obj in obj.resources.items():
            self.list_attribute(level + 1, res_name, res_obj)

    def list_attributes(self, parse_tree, select_resource=None):

        """
        list all properties (attributes) in an table.
        create the table and fill it up
        :param parse_tree:
        :param select_resource:
        """
        self.tableAttribute = self.document.add_table(rows=1, cols=5, style='TABLE-A')
        hdr_cells = self.tableAttribute.rows[0].cells
        hdr_cells[0].text = 'Property name'
        hdr_cells[1].text = 'Value type'
        hdr_cells[2].text = 'Mandatory'
        hdr_cells[3].text = 'Access mode'
        hdr_cells[4].text = 'Description'

        level = 1

        if select_resource is None:
            for resource, obj in parse_tree.resources.items():
                self.list_attribute(level, resource, obj)
        else:
            for resource, obj in parse_tree.resources.items():
                if resource[1:] == select_resource:
                    self.list_attribute(level, resource, obj)

        if self.sensor_switch is True:
            # auto generate the sensor value data..
            row_cells = self.tableAttribute.add_row().cells
            row_cells[0].text = "value"
            row_cells[1].text = "boolean"
            row_cells[2].text = "yes"
            row_cells[3].text = "Read Only"
            row_cells[4].text = "True = Sensed, False = Not Sensed."

        if self.schema_switch is True:
            # add values from external schema.
            for schema_file in self.schema_files:
                linestring = open(schema_file, 'r').read()
                # add fields in table with contents..
                self.parse_schema(linestring)

    def remove_eof_smart(self, input_string):
        """
        removes all EOL of the input string
        needed for Introduction

        :param input_string: string
        :return: string without EOL chars
        """
        ret_string = ""
        lines = input_string.splitlines()
        for line in lines:
            # add behind the added line an space
            ret_string = ret_string + line + " "
        return ret_string

    def add_justification_smart(self, depth, input_string):

        """
        add the spaces for an correct indentation of the generated RAML code section
        for descriptions in the RAML definitions
        :param depth: character depth
        :param input_string: string to be adjusted
        :return:  adjusted string
        """
        ret_string = ""
        all_lines = input_string.splitlines()
        for x_line in all_lines:
            lines = x_line.split(". ")
            for line in lines:
                string1 = depth + line + "\n"
                if len(line) > 0:
                    ret_string = ret_string + string1
        return ret_string

    def add_justification(self, depth, input_string):
        """
        add the spaces for an correct indentation of the generated RAML code section
        needed for schema and code
        :param depth: character depth
        :param input_string: string to be adjusted
        :return: adjusted string
        """
        ret_string = ""
        lines = input_string.splitlines()
        for line in lines:
            string1 = depth + line + "\n"
            ret_string = ret_string + string1
        return ret_string

    def print_bodies(self, depth, bodies):
        """
        function to loop over the bodies in an method
        :param depth: character depth
        :param bodies: object
        """
        if bodies is not None:
            for b_name, body in bodies.items():
                self.print_body(depth, b_name, body)

    def print_post_put_body(self, depth, b_name, body):
        """
        hack for parse tree in put/post
        :param depth: character depth
        :param b_name: body name
        :param body: object
        """
        tdepth = depth + self.tab
        self.document.add_paragraph(depth + "body:", style='CODE-AQUA')
        post_txt = tdepth + "application/json" + ":"
        self.document.add_paragraph(post_txt, style='CODE-AQUA')
        self.print_body(tdepth, b_name, body)

    def validate_with_json_lint(self, schema_filename, json_file, json_string):
        """
        validated the json by means of json lint

        :param schema_filename:
        :param json_file:
        :param json_string:
        """
        return
        print "validating ", json_file, " with ", schema_filename
        print "==>validate_with_json_lint: validation start:"
        my_cmd = "jsonlint " + json_file + " -V " + schema_filename
        try:
            os.system(my_cmd)
        except OSError, e:
            print >> sys.stderr, "Execution failed:", e

        print "==> validate_with_json_lint: validation complete"

    def validate_body(self, body):
        schema_string = ""
        try:
            # validation by using package:
            # https://pypi.python.org/pypi/jsonschema
            print "xx=> validation schema (jsonschema)"
            schema_string = self.get_schema_string_from_body(body)

            v_schema = None
            v_example = None
            try:
                v_schema = json.loads(schema_string)
            except ValueError as ex:
                print "error with loading schema:"
                print ex
            try:
                v_example = json.loads(body.example)
            except ValueError as ex:
                print "error with loading example:"
                print ex

            # Lazily report all errors in the instance
            validation_error = False
            v = None
            try:
                v = Draft4Validator(v_schema)
                for error in sorted(v.iter_errors(v_example), key=str):
                    validation_error = True
                    print(error.message)
                    print(error)

            except ValidationError as e:
                validation_error = True
                print "validation failed:"
                print (e.message)

                for error in sorted(v.iter_errors(v_example), key=str):
                    validation_error = True
                    print(error.message)
                    print(error)

            if validation_error is True:
                print "validation failed, input information:"
                print "body (json):"
                print body.example
                print ""
                print "schema (json):"
                print schema_string
            else:
                print "xx=xx=> schema & json VALID"
                # validate (v_example, v_schema, cls=Draft4Validator)
        except:
            print ""
            print "ERROR: failure in body (json):"
            print body.example
            print ""
            print "schema (json):"
            print schema_string
            print ""

        print "xxx=> validation schema (jsonschema) done"

    def print_body(self, depth, body_name, body):
        """

        :param depth: depth in characters
        :param body_name:
        :param body: body object to be put in the word document
        :return:
        """
        tdepth = depth + self.tab
        ttdepth = tdepth + self.tab
        tttdepth = ttdepth + self.tab
        write_depth = tdepth
        filename = ""

        if body.schema is None:
            return
        if body.schema is not None:
            if body_name is not "":
                self.document.add_paragraph(tdepth + "body:", style='CODE-AQUA')
                method_txt = ttdepth + body_name + ":"
                self.document.add_paragraph(method_txt, style='CODE-AQUA')
                write_depth = tttdepth

            p = self.document.add_paragraph(write_depth + "schema", style='CODE-GREY')
            p.add_run(": |", style='CODE-GREY')
            # schema itself
            schema_text = self.get_schema_string_from_body(body)
            try:
                schema_text = str(schema_text)
                par = self.document.add_paragraph(self.add_justification(write_depth + self.tab, schema_text),
                                                  style='CODE-BLACK')
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except:
                print "failure in (schema):", schema_text

        if body.example is not None:
            try:
                # this is a simple check if the json is correctly formatted.
                json_data = json.loads(body.example)
            except:
                print "failure in (json):", body.example

            try:
                p = self.document.add_paragraph(write_depth + "example", style='CODE-GREY')
                p.add_run(": |", style='CODE-GREY')
                par = self.document.add_paragraph(self.add_justification(write_depth + self.tab, body.example),
                                                  style='CODE-BLACK')
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except:
                print "failure in (body example):", body.example

            try:
                # check based on https://www.npmjs.com/package/jsonlint
                f = open("temp.json", 'wb')
                f.write(body.example)
                f.close()
            except:
                print "failure in validating (not executed)(body example):", body.example
            #
            # do the validation
            #
            #self.validate_with_json_lint(filename, "temp.json", body.example)
            self.validate_body(body)

    def print_header(self, depth, headers):
        """
        print header to std out

        :param depth: depth to be printed (in chars)
        :param headers: the header to be printed
        """
        try:
            for header_name, header in headers.item():
                print depth, "headername:", header_name
                print depth, "headertype:", header.type
        except:
            pass

    def print_response(self, depth, response):
        """
        prints a response object

        :param depth: depth in the tree
        :param response: RamlResponse

        """
        tdepth = depth + self.tab
        ttdepth = tdepth + self.tab
        tttdepth = ttdepth + self.tab

        if response is None:
            return

        for resName, res in response.items():
            if resName is not None:
                self.document.add_paragraph(tdepth + str(resName) + ":", style='CODE-BLUE')
            if res.description is not None:
                self.print_description(ttdepth, res.description)
            if res.schema is not None:
                p = self.document.add_paragraph(ttdepth + "schema ", style='CODE-AQUA')
                p.add_run(": |", style='CODE-YELLOW')
                par = self.document.add_paragraph(self.add_justification(tttdepth, res.schema), style='CODE-BLACK')
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if res.example is not None:
                p = self.document.add_paragraph(ttdepth + "example", style='CODE-AQUA')
                p.add_run(": |", style='CODE-YELLOW')
                par = self.document.add_paragraph(self.add_justification(tttdepth, res.example), style='CODE-BLACK')
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT

            self.print_header(tdepth, res.headers)
            self.print_bodies(tdepth, res.body)

    def print_description(self, depth, description_txt):
        """
        prints the (raml) description
        :param depth: depth in chars to print the descriptoin
        :param description_txt: text to be put in the word document
        """
        if description_txt is not None:
            self.document.add_paragraph(depth + "description: |", style='CODE-YELLOW')
            adjusted_text = self.add_justification_smart(depth + self.tab, description_txt)
            par = self.document.add_paragraph(adjusted_text, style='CODE-YELLOW')
            par.italic = True

    def list_to_array(self, input_list):
        """
        generates an raml string representation of an python list
        :param input_list: python array
        :return: string as raml string representation. example = "[ 'blah', 'blah2' ]"
        """
        my_string = "["
        for x in input_list:
            comma = ", "
            my_string = my_string + '"' + x + '"' + comma
        # remove last comman (e.g. last 2 chars)
        my_string = my_string[:-2]
        my_string += "]"
        return my_string

    def list_to_string(self, input_list):
        """
        generates an raml string representation of an python list
        :rtype : string
        :param input_list: python array ["aa", "bb"
        :return: string example "aabb"
        """
        my_string = ""
        for x in input_list:
            my_string = my_string + x
        return my_string

    def print_trait_query_parameters(self, depth, query_parameters):
        """

        :param depth:
        :param query_parameters:
        """
        tdepth = depth + self.tab
        ttdepth = tdepth + self.tab
        if query_parameters is not None:
            self.document.add_paragraph(depth + "queryParameters: ", style='CODE-AQUA')
            for query_name, query_obj in query_parameters.items():
                qtext = tdepth + query_name + ":"
                self.document.add_paragraph(qtext, style='CODE-BLUE')
                for name, q_obj in query_obj.items():
                    nametext = ""
                    if name == "enum":
                        # list as an array
                        nametext = nametext + ttdepth + name + ": " + self.list_to_array(q_obj)
                    else:
                        # list as an string
                        nametext = nametext + ttdepth + name + ": " + self.list_to_string(q_obj)
                    self.document.add_paragraph(nametext, style='CODE-BLUE')

    def print_query_parameters(self, depth, query_params):
        """
        print the query params in the RAML section

        :param depth:
        :param query_params:
        """
        tdepth = depth + self.tab
        ttdepth = tdepth + self.tab
        if query_params is not None:
            self.document.add_paragraph(depth + "queryParameters: ", style='CODE-AQUA')
            for query_name, qobj in query_params.items():
                name_text = tdepth + query_name + ":"
                self.document.add_paragraph(name_text, style='CODE-BLUE')
                if qobj.enum is not None:
                    name_text = ttdepth + "enum: " + self.list_to_string(qobj.enum)
                    self.document.add_paragraph(name_text, style='CODE-BLUE')
                if qobj.type is not None:
                    name_text = ttdepth + "type: " + self.list_to_string(qobj.type)
                    self.document.add_paragraph(name_text, style='CODE-BLUE')
                if qobj.description is not None:
                    name_text = ttdepth + "description: " + self.list_to_string(qobj.description)
                    self.document.add_paragraph(name_text, style='CODE-YELLOW')
                if qobj.required is not None:
                    if qobj.required is True:
                        name_text = ttdepth + "required: true"
                    else:
                        name_text = ttdepth + "required: false"
                    self.document.add_paragraph(name_text, style='CODE-BLUE')
                if qobj.example is not None:
                    name_text = ttdepth + "example: " + self.list_to_string(qobj.example)
                    self.document.add_paragraph(name_text, style='CODE-GREY')

    def printIS_(self, depth, is_):
        # print the is string in the RAML definition.. this on resource level
        if is_ is not None:
            my_string = depth + "is : ["
            for is_name in is_:
                my_temp = "'" + is_name + "',"
                my_string += my_temp
            my_string = my_string[:-1]
            my_string += "]"
            self.document.add_paragraph(my_string, style='CODE-BLUE')

    def print_resource(self, depth, pr_resource, obj):
        """
        print the resource in the RAML section

        :param depth:
        :param pr_resource:
        :param obj:
        :return:
        """
        tdepth = depth + self.tab
        ttdepth = tdepth + self.tab

        if obj is None:
            return

        resource_text = depth + pr_resource + ":"
        self.document.add_paragraph(resource_text, style='CODE-BLUE')
        try:
            if obj.description is not None:
                self.print_description(tdepth, obj.description)
        except:
            pass
        try:
            if obj.is_ is not None:
                self.printIS_(tdepth, obj.is_)
        except:
            pass

        if obj.methods is not None:
            for method, mobj in obj.methods.items():
                # RamlMethod
                method_txt = tdepth + method + ":"
                self.document.add_paragraph(method_txt, style='CODE-AQUA')
                # description on method level
                if mobj.description is not None:
                    self.print_description(ttdepth, mobj.description)
                # print the query parameters
                if mobj.queryParameters is not None:
                    self.print_query_parameters(ttdepth, mobj.queryParameters)
                # print the body
                if mobj.body is not None:
                    self.print_post_put_body(ttdepth, "", mobj.body)
                # print the response header of the method
                self.document.add_paragraph(ttdepth + "responses :", style='CODE-AQUA')
                # print the different responses
                self.print_response(ttdepth, mobj.responses)
            # recurse...
            depth += self.tab
            for nResName, nObj in obj.resources.items():
                self.print_resource(depth, nResName, nObj)

    def print_trait(self, depth, trait_name, obj):
        """

        :param depth:
        :param trait_name:
        :param obj:
        """
        # one extra, due to array item indicator -
        ttdepth = "   " + self.tab
        trait_string = " - " + trait_name + " :"
        self.document.add_paragraph(trait_string, style='CODE-AQUA')
        self.print_trait_query_parameters(ttdepth, obj.queryParameters)

    def print_traits(self, depth, parse_tree):
        """

        :param depth:
        :param parse_tree:
        """
        traits = parse_tree.traits
        # function to loop over the bodies in an method
        try:
            if len(traits.items()) > 0:
                self.document.add_paragraph("traits:", style='CODE-AQUA')
            # todo first trait needs a - to indicate it is an array...
            for trait_name, obj in traits.items():
                self.print_trait(self.tab, trait_name, obj)
        except:
            print "no traits found!!"
            pass

    def generate_sections(self, parse_tree, section_name=None):
        # generate the individual sections

        # just plain output
        """

        :param parse_tree:
        :param section_name:
        """
        title_name = parse_tree.title
        if section_name is not None:
            title_name = section_name
            display_name = self.get_display_name_resources(parse_tree, section_name)
            self.displayName = display_name
            print "DisplayName:", display_name
            if display_name is not None:
                title_name = display_name
        print "Title", title_name
        self.title = title_name

        print "RT = ", self.get_resource_type_by_resources(parse_tree, section_name)

        # section Resource name
        par = self.document.add_heading(title_name, level=2)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading1'
        # section introduction
        par = self.document.add_heading('Introduction', level=3)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading2'
        self.list_descriptions(parse_tree, select_resource=section_name)

        # section URI
        if self.annex_switch is False:
            par = self.document.add_heading('Example URI', level=3)
        else:
            par = self.document.add_heading('Wellknown URI', level=3)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading2'
        self.list_URIs(parse_tree, select_resource=section_name)

        # section RT
        par = self.document.add_heading('Resource Type', level=3)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading2'
        rt_name = self.get_resource_type_by_resources(parse_tree, section_name)
        if rt_name is not None:
            text = "The resource type (rt) is defined as: " + rt_name + "."
            self.document.add_paragraph(text)
        else:
            print "RT not found!"

        # section RAML definition
        par = self.document.add_heading('RAML Definition', level=3)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading2'

        # self.document.add_section()
        self.document.add_paragraph("#%RAML 0.8", style='CODE-GREEN')
        p = self.document.add_paragraph("title: ", style='CODE-YELLOW')
        p.add_run(parse_tree.title).italic = True
        p = self.document.add_paragraph("version: ", style='CODE-YELLOW')
        version_text = str(parse_tree.version)
        p.add_run(version_text).italic = True

        self.print_traits("", parse_tree)

        self.document.add_paragraph("")
        for resource, obj in parse_tree.resources.items():
            if section_name is None:
                self.print_resource("", resource, obj)
            else:
                if section_name == resource[1:]:
                    self.print_resource("", resource, obj)

        if self.composite_switch is False:
            # do not add when the switch is true...
            # section property definition
            par = self.document.add_heading('Property Definition', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'
            self.list_attributes(parse_tree, select_resource=section_name)

        # section CRUDN definition
        par = self.document.add_heading('CRUDN behavior', level=3)
        if self.annex_switch is True:
            par.style = 'ANNEX-heading2'
        self.list_resources_crudn(parse_tree, select_resource=section_name)

        if self.schema_switch is True:
            # section extra JSON definition
            par = self.document.add_heading('Referenced JSON schemas', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'

            for my_schema_file in self.schema_files:
                par = self.document.add_heading(my_schema_file, level=4)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'
                schema_text = open(my_schema_file, 'r').read()
                try:
                    par = self.document.add_paragraph(self.add_justification("", schema_text), style='CODE-BLACK')
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except:
                    pass

        if self.schemaWT_switch is True:
            # section extra JSON definition
            par = self.document.add_heading('Referenced JSON schemas', level=3)
            if self.annex_switch is True:
                par.style = 'ANNEX-heading2'

            for schema_file in self.schemaWT_files:
                par = self.document.add_heading(schema_file, level=4)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'

                par = self.document.add_heading("Property Definition", level=5)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'

                schema_text = open(schema_file, 'r').read()

                self.tableAttribute = self.document.add_table(rows=1, cols=5, style='TABLE-A')

                hdr_cells = self.tableAttribute.rows[0].cells
                hdr_cells[0].text = 'Property name'
                hdr_cells[1].text = 'Value type'
                hdr_cells[2].text = 'Mandatory'
                hdr_cells[3].text = 'Access mode'
                hdr_cells[4].text = 'Description'

                # add fields in table with contents..
                self.parse_schema(schema_text)
                par = self.document.add_heading("Schema Definition", level=5)
                if self.annex_switch is True:
                    par.style = 'ANNEX-heading2'
                try:
                    par = self.document.add_paragraph(self.add_justification("", schema_text), style='CODE-BLACK')
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except:
                    pass

    def schemaRef2Filename(self, schema_name):
        # convert the schema reference name into the actual filename to be read
        schemas = self.parsetree.schemas
        try:
            for item in schemas:
                for name, obj in item.items():
                    if name == schema_name:
                        return obj.file_name
        except:
            pass
        return "ERROR-IN-RESOLVING-SCHEMA:NO_FILE_FOUND_FOR:" + str(schema_name)

    def get_schema_string_from_body(self, body):
        """
        convert the schema reference name into the actual filename to be read
        will determine if the referenced file needs to be read..

        :param body:
        :return:
        """
        schema_string = None
        try:
            schema_string = str(body.schema)
        except:
            pass
        if schema_string is not None:
            if "{" not in schema_string:
                # we think this is a reference.
                # find it and include it.
                filename = self.schemaRef2Filename(schema_string)
                print "resolve schema reference:", schema_string, filename
                # read the file as a string
                try:
                    schema_string = self.read_file(filename)
                except:
                    print "could not open file:", filename
        #print "get_schema_string_from_body", schema_string
        return schema_string

    def read_file(self, filename):
        """
        read the file as a string

        :param filename: file to read
        :return:
        """
        try:
            linestring = open(filename, 'r').read()
            # create the table with contents..
            return linestring
        except:
            pass
        try:
            full_path = os.path.join(self.dir, filename)
            linestring = open(full_path, 'r').read()
            # create the table with contents..
            return linestring
        except:
            pass

        try:
            base = os.path.basename(filename)
            full_path = os.path.join(self.dir, base)
            linestring = open(full_path, 'r').read()
            # create the table with contents..
            return linestring
        except:
            pass

        print "read_file: could not open file:", filename, full_path

    def convert(self):
        """
        conversion of the raml info into the word document

        :return:
        """
        try:
            parsetree = ramlparser.load(self.inputname)
        except ValidationError as e:
            print 'validation error:', e.errors
            print "could not load file: error loading file"
            traceback.print_exc()
            return

        # make it a member..
        self.parsetree = parsetree

        self.list_x_resources(parsetree)

        # print parsetree
        # output = dump(parsetree, Dumper=Dumper,default_flow_style=False)
        # output = dump(parsetree, Dumper=SafeDumper)
        # print output

        try:
            self.document = Document(docx=self.resourcedoc)
        except:
            print "could not load file: ", self.resourcedoc
            print "make sure that docx file exist.."
            return

        self.generate_sections(parsetree, self.resource_name)

        self.document.save(self.resource_out)
        print "document saved..", self.resource_out

    def add_header(self, level_nr, header_title):
        """
        add an header to the document.
        "_" are replaced to " ", underscores are replaced by chars.

        :param level_nr:
        :param header_title:
        :return:
        """
        try:
            self.document = Document(docx=self.resourcedoc)
        except:
            print "could not load file: ", self.resourcedoc
            print "make sure that docx file exist.."
            return

        header = str(header_title).replace("_", " ")
        print "add_header: title:", header
        paragraph = self.document.add_heading(header, level=1)
        if self.annex_switch is True:
            print "   as annex"
            paragraph.style = 'ANNEX_title'

        self.document.save(self.resource_out)


#
# code for the proxy
#
import os
import SimpleHTTPServer
import SocketServer
import threading
import urllib2


class ProxyHandler(SimpleHTTPServer.SimpleHTTPRequestHandler):
    def do_GET(self):
        """
        serve up an file
        """

        schema_dir = args['schemadir']
        #raml_file = args['ramlName']

        # filename without any path
        filename = self.path.split('/')[-1]
        base_name = os.path.basename(filename)
        print "ProxyHandler: url:", self.path, " localfile:", base_name
        if os.path.exists(base_name):
            print "ProxyHandler: local file found:", base_name
            self.copyfile(open(base_name), self.wfile)
            return

        # filename with path of schemas
        full_path = os.path.join(schema_dir, base_name)
        print "ProxyHandler: url:", self.path, " localfile:", full_path
        if os.path.exists(full_path):
            print "ProxyHandler: local file found:", full_path
            self.copyfile(open(full_path), self.wfile)
            return

        filenamejson = base_name + ".json"
        print "ProxyHandler: local file NOT found:", base_name, " trying: ", filenamejson
        if os.path.exists(filenamejson):
            print "ProxyHandler: local file found:", filenamejson
            self.copyfile(open(filenamejson), self.wfile)
        else:

            print "ProxyHandler: trying url:", self.path
            proxy_handler = urllib2.ProxyHandler({})
            opener = urllib2.build_opener(proxy_handler)
            try:
                req = urllib2.Request(self.path)
                self.copyfile(opener.open(req), self.wfile)
            except:
                print "ProxyHandler: file not found:", self.path


def proxy():
    """
    start proxy for serving up the referenced schema files

    :rtype : null
    """
    PORT = 4321
    httpd = SocketServer.TCPServer(("", PORT), ProxyHandler)
    print "HTTPPRoxy: serving at port", PORT
    proxythread = threading.Thread(target=httpd.serve_forever)
    proxythread.setDaemon(True)
    proxythread.start()
    os.environ['http_proxy'] = 'http://localhost:%d/' % PORT
    return


if __name__ == '__main__':

    resourcedoc = "ResourceTemplate.docx"
    # set the execution path of the tool
    if hasattr(sys, 'frozen'):
        my_dir = os.path.dirname(sys.executable)
    else:
        my_dir = os.path.dirname(sys.argv[0])


    # version information
    my_version = ""
    try:
        from version import VERSION

        my_version = VERSION
    except:
        pass

    print "==================================="
    print "version: ", my_version

    # annex_switch = False

    # argument parsing
    parser = argparse.ArgumentParser(description='Process RAML files.')
    parser.add_argument('-docx', '--docx', help='word template file')
    parser.add_argument('-outdocx', '--outdocx', help='word output file')
    parser.add_argument('-raml', '--raml', help='raml input file')
    parser.add_argument('-schemadir', '--schemadir', help='schema dir input file')
    parser.add_argument('-heading1', '--heading1', help='creates an heading 1 to the document (and exit)')
    # parser.add_option('-showResources','--showResources', help='shows the resources in an RAML file')
    parser.add_argument('-resource', '--resource', help='resource to be processed')

    parser.add_argument('-annex', '--annex', help='uses a annex heading instead of normal heading (--annex true)')
    parser.add_argument('-put', '--put', help='uses put command as property table input instead of get (--put true)')
    parser.add_argument('-composite', '--composite',
         help='treats the resource as an composite resource, e.g. no property definition table (--composite true)')
    parser.add_argument('-sensor', '--sensor',
         help='treats the resource as an sensor resource, e.g. add the value "value" to the property table (--sensor true)')
    parser.add_argument('-schema', '--schema', nargs='*',
         help='additional (referenced) schema used in the resource (--schema "schema file1" "schema file2" )')
    parser.add_argument('-schemaWT', '--schemaWT', nargs='*',
         help='additional (referenced) schema (section With Table) used in the resource (--schema "schema file1" "schema file2" )')

    args = vars(parser.parse_args())

    resourceName = args['resource']
    docxName = args['docx']
    ramlName = args['raml']
    header0 = args['heading1']
    annex_switch = args['annex']
    put_switch = args['put']
    composite_switch = args['composite']
    sensor_switch = args['sensor']
    schema_file = args['schema']
    schemaWT_file = args['schemaWT']

    # annex_switch = True

    if annex_switch is None:
        annex_switch = False
    else:
        annex_switch = True

    if put_switch is None:
        put_switch = False
        table_method = 'get'
    else:
        put_switch = True
        table_method = 'put'

    if composite_switch is None:
        composite_switch = False
    else:
        composite_switch = True

    if sensor_switch is None:
        sensor_switch = False
    else:
        sensor_switch = True

    if schema_file is None:
        # schema_file to be included in class
        schema_switch = False
    else:
        schema_switch = True

    if schemaWT_file is None:
        # schema_file to be included in class
        schemaWT_switch = False
    else:
        schemaWT_switch = True

    if docxName is None:
        docxName = resourcedoc

    print "==================================="
    print "using raml file              :", ramlName
    print "using docx file              :", docxName
    print "using docx output file       :", args['outdocx']
    print "using schema dir             :", args['schemadir']
    print "using resource               :", resourceName
    print "using header0                :", header0
    print "using annex                  :", annex_switch
    print "using put for property table :", put_switch
    print "using composite              :", composite_switch
    print "using sensor                 :", sensor_switch
    print "schema switch                :", schema_switch
    print "schema (WT) switch           :", schemaWT_switch
    if schema_switch == True:
        print "schema file                  :", schema_file
    if schemaWT_switch == True:
        print "schema (WT) file             :", schemaWT_file

    print "styles:"
    print " heading: Heading 1 or ANNEX-heading1"
    print " table style: TABLE-A"
    print " table header style: TABLEHEADER"
    print " color (code) style: CODE-AQUA"
    print "                   : CODE-YELLOW"
    print "                   : CODE-GREY"
    print "                   : CODE-BLACK"
    print "                   : CODE-BLUE"
    print "                   : CODE-GREEN"
    print "==================================="

    temp = sys.stdout
    # sys.stdout = sys.stderr
    sys.stderr = sys.stdout

    # start the proxy
    proxy()

    if my_dir:
        os.chdir(my_dir)

    if len(sys.argv) == 1:
        parser.print_help()
        processor = None
    else:
        processor = CreateDoc(ramlName, docx_name=docxName, resource_name=resourceName)

    if processor is not None:
        processor.annex_switch = annex_switch
        processor.composite_switch = composite_switch
        processor.sensor_switch = sensor_switch
        processor.table_method = table_method
        processor.sensor_switch = sensor_switch
        processor.schema_switch = schema_switch
        processor.schemaWT_switch = schemaWT_switch
        processor.dir = args['schemadir']
        if args['outdocx'] is not None:
            processor.resource_out = args['outdocx']
        if schema_switch is True:
            processor.schema_files = schema_file
        if schemaWT_switch is True:
            processor.schemaWT_files = schemaWT_file
        if header0 is not None:
            processor.add_header(0, header0)
            sys.exit()

    if processor is not None:
        processor.convert()

    for resource, obj in processor.parsetree.resources.items():
        print "resource :", resource